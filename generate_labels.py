#!/usr/bin/env python3
"""
generate_labels.py
-------------------
Converts an ERP (EasyGems) sales-invoice Excel export directly into a
print-ready Word label sheet -- replicating the exact layout used in
STICKER_NEW.docx (5 labels across x 8 rows down = 40 labels per A4 page,
custom laser label stock: 39mm wide x 35mm tall).

Usage:
    python3 generate_labels.py <input_excel.xlsx> <output_labels.docx> [--sheet SHEETNAME]

No manual column deletion, mail-merge setup, or formatting is required --
point it at the raw ERP export and it produces the finished, ready-to-print
Word file.
"""

import sys
import argparse
import openpyxl
from docx import Document
from docx.shared import Twips, Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------
# Layout constants (lifted directly from the proven STICKER_NEW.docx)
# ----------------------------------------------------------------------
PAGE_WIDTH_TW = 11906      # A4
PAGE_HEIGHT_TW = 16838
MARGIN_TOP_TW = 170
MARGIN_RIGHT_TW = 255
MARGIN_BOTTOM_TW = 0
MARGIN_LEFT_TW = 255

COLS = 5
ROWS = 8
PER_PAGE = COLS * ROWS      # 40 labels per page
COL_WIDTH_TW = 2296
ROW_HEIGHT_TW = 2070
CELL_MARGIN_LR_TW = 43

FONT_NAME = "Calibri"
BODY_SIZE_PT = 12
MM_LINE_SIZE_PT = 11

SPACER_AFTER_SSP = " " * 15   # between SSP code and Item#
SPACER_BEFORE_LAB = " " * 7   # indent before "Lab - Report#" line


# ----------------------------------------------------------------------
# Header matching -- tolerant of the exact column names an ERP export
# happens to use (raw EasyGems names, or the "cleaned" friendly names
# described in the manual workflow).
# ----------------------------------------------------------------------
CANDIDATES = {
    "ssp":      ["ssp", "fssp", "ssp any", "lot #", "lot#"],
    "item":     ["item#", "item #", "item", "item number"],
    "stntyp":   ["stntyp", "stone type", "gem"],
    "shp":      ["shp", "shape"],
    "col":      ["col", "color", "colour"],
    "units":    ["units", "wt", "weight", "carat", "carats", "cts"],
    "lab_report_combined": ["lab report number", "lab report#", "lab report #",
                             "lab-report", "labreport"],
    "lab":      ["lab"],
    "report":   ["report#", "report #", "report number", "report"],
    "length":   ["l", "length"],
    "width":    ["w", "width"],
    "depth":    ["d", "depth"],
}


def normalize(h):
    return str(h).strip().lower() if h is not None else ""


def build_header_map(header_row):
    norm_to_orig = {}
    for idx, h in enumerate(header_row):
        n = normalize(h)
        if n and n not in norm_to_orig:
            norm_to_orig[n] = idx
    field_idx = {}
    for field, names in CANDIDATES.items():
        for name in names:
            if name in norm_to_orig:
                field_idx[field] = norm_to_orig[name]
                break
    return field_idx


def fmt_num(v):
    """General-number formatting: strip unnecessary decimal zeros,
    matching how the source ERP values print (e.g. 8.0 -> '8',
    11.90 -> '11.9', 12.38 -> '12.38')."""
    if v is None or v == "":
        return None
    try:
        f = float(v)
    except (ValueError, TypeError):
        return str(v)
    if f == int(f):
        return str(int(f))
    s = f"{f:.2f}".rstrip("0").rstrip(".")
    return s


def build_records(ws, field_idx):
    records = []
    rows = ws.iter_rows(min_row=2, values_only=True)
    for row in rows:
        def get(field):
            idx = field_idx.get(field)
            return row[idx] if idx is not None and idx < len(row) else None

        ssp = get("ssp")
        item = get("item")
        stntyp = get("stntyp")
        shp = get("shp")
        col = get("col")
        units = get("units")
        length = get("length")
        width = get("width")
        depth = get("depth")

        # Skip completely empty rows
        if all(v in (None, "") for v in
               [ssp, item, stntyp, shp, col, units, length, width, depth]):
            continue

        # Lab - Report# line
        combined = get("lab_report_combined")
        if combined not in (None, ""):
            f10 = str(combined)
        else:
            lab = get("lab")
            report = get("report")
            if lab not in (None, "") and report not in (None, ""):
                f10 = f"{lab} - {report}"
            else:
                f10 = ""

        # Weight line (always 2 decimals, matches source convention)
        if units not in (None, ""):
            try:
                f6 = f"{float(units):.2f} Cts"
            except (ValueError, TypeError):
                f6 = f"{units} Cts"
        else:
            f6 = ""

        # MM line (general formatting, blank if no real dimensions recorded)
        def nz(v):
            try:
                return v is not None and float(v) != 0
            except (ValueError, TypeError):
                return v not in (None, "")

        if nz(length) or nz(width) or nz(depth):
            l_s, w_s, d_s = fmt_num(length) or "0", fmt_num(width) or "0", fmt_num(depth) or "0"
            f14 = f"{l_s} - {w_s} X {d_s}"
        else:
            f14 = ""

        records.append({
            "ssp": ssp if ssp is not None else "",
            "item": item if item is not None else "",
            "f10": f10,
            "stntyp": stntyp if stntyp is not None else "",
            "shp": shp if shp is not None else "",
            "col": col if col is not None else "",
            "f6": f6,
            "f14": f14,
        })
    return records


# ----------------------------------------------------------------------
# DOCX construction helpers
# ----------------------------------------------------------------------
def set_section_geometry(section):
    section.page_width = Twips(PAGE_WIDTH_TW)
    section.page_height = Twips(PAGE_HEIGHT_TW)
    section.top_margin = Twips(MARGIN_TOP_TW)
    section.right_margin = Twips(MARGIN_RIGHT_TW)
    section.bottom_margin = Twips(MARGIN_BOTTOM_TW)
    section.left_margin = Twips(MARGIN_LEFT_TW)


def set_table_fixed_no_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    tbl_look = tblPr.find(qn("w:tblLook"))
    insert_idx = list(tblPr).index(tbl_look) if tbl_look is not None else len(list(tblPr))

    elements = []

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    elements.append(borders)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    elements.append(layout)

    cellmar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", 0), ("bottom", 0),
                      ("left", CELL_MARGIN_LR_TW), ("right", CELL_MARGIN_LR_TW)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        cellmar.append(el)
    elements.append(cellmar)

    for offset, el in enumerate(elements):
        tblPr.insert(insert_idx + offset, el)


def set_exact_row_height(row, height_tw):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)
    height = OxmlElement("w:trHeight")
    height.set(qn("w:hRule"), "exact")
    height.set(qn("w:val"), str(height_tw))
    trPr.append(height)


def set_col_widths(table):
    table.autofit = False
    for col in table.columns:
        col.width = Twips(COL_WIDTH_TW)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Twips(COL_WIDTH_TW)


def add_label_paragraph(cell, text, size_pt=BODY_SIZE_PT, first=False):
    p = cell.add_paragraph() if not first else cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    return p


def fill_label_cell(cell, rec):
    line1 = f"{rec['ssp']}{SPACER_AFTER_SSP}{rec['item']}"
    line2 = f"{SPACER_BEFORE_LAB}{rec['f10']}"
    line3 = f"Gem \u2013 {rec['stntyp']}"
    line4 = f"Shape \u2013 {rec['shp']}"
    line5 = f"Col \u2013 {rec['col']}"
    line6 = f"Wt \u2013 {rec['f6']}"
    line7 = f"MM \u2013 {rec['f14']}"

    add_label_paragraph(cell, line1, first=True)
    add_label_paragraph(cell, line2)
    add_label_paragraph(cell, line3)
    add_label_paragraph(cell, line4)
    add_label_paragraph(cell, line5)
    add_label_paragraph(cell, line6)
    add_label_paragraph(cell, line7, size_pt=MM_LINE_SIZE_PT)
    add_label_paragraph(cell, "")  # trailing spacer line


def build_document(records, out_path):
    doc = Document()
    set_section_geometry(doc.sections[0])

    chunks = [records[i:i + PER_PAGE] for i in range(0, len(records), PER_PAGE)] or [[]]

    for page_idx, chunk in enumerate(chunks):
        table = doc.add_table(rows=ROWS, cols=COLS)
        set_table_fixed_no_border(table)
        set_col_widths(table)
        for row in table.rows:
            set_exact_row_height(row, ROW_HEIGHT_TW)

        for i, rec in enumerate(chunk):
            r, c = divmod(i, COLS)
            fill_label_cell(table.cell(r, c), rec)

        if page_idx < len(chunks) - 1:
            doc.add_page_break()

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate a printable label sheet from an ERP Excel export.")
    parser.add_argument("input_excel")
    parser.add_argument("output_docx")
    parser.add_argument("--sheet", default=None, help="Sheet name (defaults to active/first sheet)")
    args = parser.parse_args()

    wb = openpyxl.load_workbook(args.input_excel, data_only=True)
    ws = wb[args.sheet] if args.sheet else wb.active

    header_row = [c.value for c in ws[1]]
    field_idx = build_header_map(header_row)

    required = ["ssp", "item", "stntyp", "shp", "col", "units"]
    missing = [f for f in required if f not in field_idx]
    if missing:
        print(f"WARNING: could not find columns for: {missing}. "
              f"These fields will be left blank on the labels.", file=sys.stderr)

    records = build_records(ws, field_idx)
    if not records:
        print("ERROR: no data rows found in the sheet.", file=sys.stderr)
        sys.exit(1)

    build_document(records, args.output_docx)
    print(f"Generated {len(records)} labels across "
          f"{(len(records) + PER_PAGE - 1) // PER_PAGE} page(s) -> {args.output_docx}")


if __name__ == "__main__":
    main()
